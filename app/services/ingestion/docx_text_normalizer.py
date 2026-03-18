from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument

from app.services.ingestion.document_ingestion_pipeline import (
    NormalizationInput,
    NormalizationResult,
)


class DocxTextNormalizer:
    """
    DOCX normalizer compatible with current DocumentIngestionPipeline.

    Responsibilities:
    - read .docx from disk
    - extract plain text conservatively
    - normalize whitespace
    - compute normalized content hash
    - return NormalizationResult expected by the pipeline

    Notes:
    - this layer is intentionally text-only
    - table-aware structuring is handled later by DocxStructureExtractor
    """

    async def normalize(self, payload: NormalizationInput) -> NormalizationResult:
        file_path = Path(payload.file_path)

        if file_path.suffix.lower() != ".docx":
            raise ValueError(
                f"DocxTextNormalizer supports only .docx files, got: {file_path.suffix}"
            )

        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {payload.file_path}")

        document = DocxDocument(str(file_path))

        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = self._clean_text(paragraph.text)
            if text:
                parts.append(text)

        # Conservative fallback:
        # include flattened table text in normalized_text so downstream
        # metadata detection can still see table content if needed.
        for table in document.tables:
            for row in table.rows:
                row_values = [self._clean_text(cell.text) for cell in row.cells]
                row_values = [value for value in row_values if value]
                if row_values:
                    parts.append(" | ".join(row_values))

        normalized_text = self._normalize_whitespace("\n".join(parts))
        normalized_content_hash = self._sha256(normalized_text)

        parser_payload_json = {
            "normalizer": "docx_text_normalizer",
            "source_format": "docx",
            "paragraphs_count": len(document.paragraphs),
            "tables_count": len(document.tables),
            "detected_extension": file_path.suffix.lower(),
            "source_type": payload.source_type,
        }

        return NormalizationResult(
            normalized_text=normalized_text,
            normalized_content_hash=normalized_content_hash,
            detected_language_code=self._detect_language_code(normalized_text),
            parser_payload_json=parser_payload_json,
        )

    def _clean_text(self, value: str) -> str:
        text = value.replace("\xa0", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\s*\n\s*", "\n", text)
        return text.strip()

    def _normalize_whitespace(self, value: str) -> str:
        text = value.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()

    def _sha256(self, value: str) -> str:
        return hashlib.sha256(value.encode("utf-8")).hexdigest()

    def _detect_language_code(self, text: str) -> Optional[str]:
        if not text:
            return None

        cyrillic_chars = len(re.findall(r"[Р-пр-џЈИ]", text))
        latin_chars = len(re.findall(r"[A-Za-z]", text))

        if cyrillic_chars == 0 and latin_chars == 0:
            return None

        if cyrillic_chars >= latin_chars:
            return "ru"

        return "en"