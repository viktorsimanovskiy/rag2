from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
import tempfile
import zipfile
from collections import Counter
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable, Optional

from docx import Document as DocxDocument
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

CORE_TABLE_TYPES = {"identifiers", "documents", "refusal_reasons"}
OFFICIAL_START_MARKER = "МИНИСТЕРСТВО СОЦИАЛЬНОЙ ПОЛИТИКИ"

CONSULTANT_NOISE_MARKERS = (
    "КОНСУЛЬТАНТПЛЮС",
    "ДОКУМЕНТ ПРЕДОСТАВЛЕН КОНСУЛЬТАНТПЛЮС",
    "WWW.CONSULTANT.RU",
    "СПИСОК ИЗМЕНЯЮЩИХ ДОКУМЕНТОВ",
    "НУМЕРАЦИЯ ПУНКТОВ ДАНА В СООТВЕТСТВИИ",
)

FORM_OR_CONSENT_MARKERS = (
    "ЗАЯВЛЕНИЕ",
    "СОГЛАСИЕ",
    "РАСПИСК",
    "УВЕДОМЛЕНИЕ",
    "ПОДПИСЬ ЗАЯВИТЕЛ",
    "ДОСТОВЕРНОСТЬ И ПОЛНОТУ СВЕДЕНИЙ",
    "ПРИНЯЛ ДОКУМЕНТЫ",
    "ПРИЛАГАЮ СЛЕДУЮЩИЕ ДОКУМЕНТЫ",
    "К ЗАЯВЛЕНИЮ ПРИЛАГАЮ",
    "ПЕРЕЧЕНЬ ПРИЛАГАЕМЫХ ДОКУМЕНТОВ",
    "СВЕДЕНИЯ О ДОКУМЕНТЕ",
    "СВЕДЕНИЯ О ДОКУМЕНТАХ",
    "СВЕДЕНИЯ О ПРЕДСТАВИТЕЛЕ",
    "СВЕДЕНИЯ О ЗАКОННОМ ПРЕДСТАВИТЕЛЕ",
    "БАНКОВСКИЕ РЕКВИЗИТЫ",
    "ИНДИВИДУАЛЬНЫЙ ЛИЦЕВОЙ СЧЕТ",
)

PROCEDURE_OR_SCHEME_MARKERS = (
    "БЛОК-СХЕМ",
    "АДМИНИСТРАТИВНЫХ ПРОЦЕДУР",
    "АДМИНИСТРАТИВНЫЕ ПРОЦЕДУРЫ",
)


@dataclass(slots=True)
class OrderedItem:
    seq: int
    kind: str
    text: str
    table_index: Optional[int] = None
    rows_count: Optional[int] = None


@dataclass(slots=True)
class DocumentAnalysis:
    cleaned_filename: str
    raw_filename: Optional[str]
    raw_found: bool
    raw_size_bytes: Optional[int]
    raw_sha256: Optional[str]
    cleaned_size_bytes: int
    cleaned_sha256: str
    raw_official_start_seq: Optional[int]
    raw_items_before_official_start: Optional[int]
    cleaned_official_start_seq: Optional[int]
    cleaned_items_before_official_start: Optional[int]
    cleaned_tables_count: int
    core_table_counts: dict[str, int]
    core_table_indexes: dict[str, list[int]]
    has_all_core_tables: bool
    has_exactly_one_each_core_table: bool
    last_core_table_index: Optional[int]
    last_core_table_type: Optional[str]
    last_core_table_seq: Optional[int]
    trim_candidate: bool
    tail_items_count: int
    tail_paragraphs_count: int
    tail_tables_count: int
    tail_chars_count: int
    tail_table_type_counts: dict[str, int]
    tail_category_counts: dict[str, int]
    tail_contains_core_table: bool
    tail_first_samples: list[dict[str, Any]]
    remaining_consultant_noise_items: int
    warnings: list[str]
    error: Optional[str] = None


@dataclass(slots=True)
class RunSummary:
    analyzed_files: int
    errors_count: int
    raw_matched_files: int
    raw_official_start_found: int
    cleaned_official_start_found: int
    exact_core_1_1_1_files: int
    all_core_found_files: int
    trim_candidate_files: int
    tail_items_total: int
    tail_paragraphs_total: int
    tail_tables_total: int
    tail_chars_total: int
    remaining_consultant_noise_items_total: int
    tail_contains_core_table_files: int
    core_count_patterns: dict[str, int]
    tail_table_type_counts: dict[str, int]
    tail_category_counts: dict[str, int]


def _clean_text(value: str | None) -> str:
    if not value:
        return ""
    text = value.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _norm(value: str | None) -> str:
    return _clean_text(value).replace("ё", "е").replace("Ё", "Е").upper()


def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _safe_extract_zip(zip_path: Path, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path) as archive:
        for member in archive.infolist():
            member_path = output_dir / member.filename
            resolved = member_path.resolve()
            if not str(resolved).startswith(str(output_dir.resolve())):
                raise ValueError(f"Unsafe zip member path: {member.filename}")
            if member.is_dir():
                resolved.mkdir(parents=True, exist_ok=True)
                continue
            resolved.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(member) as src, resolved.open("wb") as dst:
                shutil.copyfileobj(src, dst)


def _collect_docx_files(path: Path) -> list[Path]:
    return sorted(
        p for p in path.rglob("*.docx")
        if p.is_file() and not p.name.startswith("~$")
    )


def _strip_cleaned_suffix(filename: str) -> str:
    path = Path(filename)
    stem = path.stem
    if stem.lower().endswith("_cleaned"):
        stem = stem[: -len("_cleaned")]
    return f"{stem}{path.suffix}"


def _build_raw_map(raw_files: Iterable[Path]) -> dict[str, Path]:
    result: dict[str, Path] = {}
    for path in raw_files:
        result.setdefault(path.name.casefold(), path)
    return result


def _iter_block_items(doc: _Document) -> list[OrderedItem]:
    items: list[OrderedItem] = []
    table_index = 0
    seq = 0

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            seq += 1
            paragraph = Paragraph(child, doc)
            text = _clean_text(paragraph.text)
            if text:
                items.append(OrderedItem(seq=seq, kind="paragraph", text=text))
        elif isinstance(child, CT_Tbl):
            seq += 1
            table_index += 1
            table = Table(child, doc)
            row_texts: list[str] = []
            for row in table.rows[:14]:
                cells = [_clean_text(cell.text) for cell in row.cells]
                deduped: list[str] = []
                for cell_text in cells:
                    if cell_text and (not deduped or deduped[-1] != cell_text):
                        deduped.append(cell_text)
                if deduped:
                    row_texts.append(" | ".join(deduped))
            items.append(
                OrderedItem(
                    seq=seq,
                    kind="table",
                    table_index=table_index,
                    rows_count=len(table.rows),
                    text=" || ".join(row_texts),
                )
            )
    return items


def _official_start_seq(items: list[OrderedItem]) -> Optional[int]:
    for item in items:
        text_n = _norm(item.text)
        if text_n == OFFICIAL_START_MARKER or text_n.startswith(OFFICIAL_START_MARKER):
            return item.seq
    return None


def _count_items_before_seq(items: list[OrderedItem], seq: Optional[int]) -> Optional[int]:
    if seq is None:
        return None
    return sum(1 for item in items if item.seq < seq)


def _table_context_by_index(items: list[OrderedItem], table_index: int, context_paragraphs: int = 16) -> str:
    table_seq = next((item.seq for item in items if item.kind == "table" and item.table_index == table_index), None)
    if table_seq is None:
        return ""
    paragraphs = [item.text for item in items if item.kind == "paragraph" and item.seq < table_seq]
    return " ".join(paragraphs[-context_paragraphs:])


def _classify_table(table_item: OrderedItem, context_text: str) -> str:
    table_n = _norm(table_item.text)
    # Берём хвост контекста, чтобы заголовок предыдущей ключевой таблицы не заражал формы ниже.
    context_n = _norm(" ".join(_clean_text(context_text).split()[-260:]))
    combined_n = f"{context_n} {table_n}"

    if any(marker in combined_n for marker in CONSULTANT_NOISE_MARKERS):
        return "consultant_noise"

    refusal_context = (
        "ИСЧЕРПЫВАЮЩИЙ ПЕРЕЧЕНЬ ОСНОВАНИЙ" in context_n
        or "ПЕРЕЧЕНЬ ОСНОВАНИЙ" in context_n
        or "ТАБЛИЦА 3" in context_n
    )
    refusal_header = (
        "НАИМЕНОВАНИЕ ОСНОВАНИЯ" in table_n
        or "ОСНОВАНИЯ ДЛЯ ОТКАЗА" in table_n
        or "ОТКАЗА В ПРИЕМЕ" in table_n
        or "ОТКАЗ В ПРИЕМЕ" in table_n
        or "ПРИОСТАНОВЛЕНИ" in table_n
    )
    if refusal_context and (refusal_header or "ОТКАЗ" in table_n or "ПРИОСТАНОВ" in table_n):
        return "refusal_reasons"
    if refusal_header and "НАИМЕНОВАН" in table_n and ("ОТКАЗ" in table_n or "ПРИОСТАНОВ" in table_n):
        return "refusal_reasons"

    documents_context = (
        "ИСЧЕРПЫВАЮЩИЙ ПЕРЕЧЕНЬ ДОКУМЕНТОВ" in context_n
        or "ПЕРЕЧЕНЬ ДОКУМЕНТОВ" in context_n
        or "ТАБЛИЦА 2" in context_n
    )
    documents_header = (
        "НАИМЕНОВАНИЕ ДОКУМЕНТА" in table_n
        or "ПЕРЕЧЕНЬ НЕОБХОДИМЫХ" in table_n
        or "ПЕРЕЧЕНЬ НЕОБХОДИМЫХ ДЛЯ ПРЕДОСТАВЛЕНИЯ" in table_n
        or "ДОКУМЕНТЫ, НЕОБХОДИМЫЕ ДЛЯ ПРЕДОСТАВЛЕНИЯ" in table_n
        or "ДОКУМЕНТОВ, НЕОБХОДИМЫХ ДЛЯ ПРЕДОСТАВЛЕНИЯ" in table_n
    )
    # В формах тоже часто есть "N | Наименование документа | Количество".
    # Поэтому без контекста ключевой таблицы не считаем это таблицей документов.
    if documents_context and documents_header:
        return "documents"

    identifiers_context = (
        "ИДЕНТИФИКАТОРЫ КАТЕГОРИЙ" in context_n
        or "КАТЕГОРИЙ (ПРИЗНАКОВ) ЗАЯВИТЕЛ" in context_n
        or "ТАБЛИЦА 1" in context_n
    )
    identifiers_header = (
        "НАИМЕНОВАНИЕ ПРИЗНАКА ЗАЯВИТЕЛ" in table_n
        or "НАИМЕНОВАНИЕ ОТДЕЛЬНЫХ ПРИЗНАКОВ ЗАЯВИТЕЛ" in table_n
        or "НАИМЕНОВАНИЕ (ЗНАЧЕНИЕ) ПРИЗНАКОВ ЗАЯВИТЕЛ" in table_n
        or "ПРИЗНАКОВ ЗАЯВИТЕЛ" in table_n
        or "ПРИЗНАКОВ) ЗАЯВИТЕЛ" in table_n
        or ("ИДЕНТИФИКАТОР" in table_n and "ЗАЯВИТЕЛ" in table_n and "КАТЕГОР" in table_n)
    )
    if identifiers_context and identifiers_header:
        return "identifiers"
    if identifiers_header and "ПРИЗНАК" in table_n:
        return "identifiers"

    if any(marker in combined_n for marker in FORM_OR_CONSENT_MARKERS):
        return "form_fields"
    return "generic"


def _item_category(item: OrderedItem) -> str:
    text_n = _norm(item.text)
    if any(marker in text_n for marker in CONSULTANT_NOISE_MARKERS):
        return "consultant_noise"
    if any(marker in text_n for marker in FORM_OR_CONSENT_MARKERS):
        return "form_or_template_table" if item.kind == "table" else "form_or_template_text"
    if any(marker in text_n for marker in PROCEDURE_OR_SCHEME_MARKERS):
        return "procedure_or_scheme_table" if item.kind == "table" else "procedure_or_scheme_text"
    if text_n.startswith("ПРИЛОЖЕНИЕ") and len(text_n) <= 280:
        return "appendix_heading"
    return "other_table" if item.kind == "table" else "other_text"


def _compact_sample(text: str, limit: int = 420) -> str:
    text = _clean_text(text)
    if len(text) <= limit:
        return text
    return f"{text[:limit].rstrip()}..."


def _analyze_one(cleaned_path: Path, raw_map: dict[str, Path], max_samples: int) -> DocumentAnalysis:
    expected_raw_name = _strip_cleaned_suffix(cleaned_path.name)
    raw_path = raw_map.get(expected_raw_name.casefold())
    warnings: list[str] = []

    raw_size_bytes: Optional[int] = None
    raw_sha256: Optional[str] = None
    raw_official_start: Optional[int] = None
    raw_before_start: Optional[int] = None

    if raw_path is not None:
        raw_size_bytes = raw_path.stat().st_size
        raw_sha256 = _sha256_file(raw_path)
        try:
            raw_items = _iter_block_items(DocxDocument(str(raw_path)))
            raw_official_start = _official_start_seq(raw_items)
            raw_before_start = _count_items_before_seq(raw_items, raw_official_start)
            if raw_official_start is None:
                warnings.append("raw_official_start_not_found")
        except Exception as exc:
            warnings.append(f"raw_read_error: {exc!r}")
    else:
        warnings.append("raw_file_not_found")

    cleaned_size_bytes = cleaned_path.stat().st_size
    cleaned_sha256 = _sha256_file(cleaned_path)

    try:
        cleaned_doc = DocxDocument(str(cleaned_path))
        items = _iter_block_items(cleaned_doc)
        cleaned_official_start = _official_start_seq(items)
        cleaned_before_start = _count_items_before_seq(items, cleaned_official_start)
        if cleaned_official_start is None:
            warnings.append("cleaned_official_start_not_found")

        table_type_by_index: dict[int, str] = {}
        for item in items:
            if item.kind != "table" or item.table_index is None:
                continue
            context = _table_context_by_index(items, item.table_index)
            table_type_by_index[item.table_index] = _classify_table(item, context)

        core_table_counts_counter: Counter[str] = Counter()
        core_table_indexes: dict[str, list[int]] = {table_type: [] for table_type in sorted(CORE_TABLE_TYPES)}
        for table_index, table_type in table_type_by_index.items():
            if table_type in CORE_TABLE_TYPES:
                core_table_counts_counter[table_type] += 1
                core_table_indexes[table_type].append(table_index)

        core_table_counts = {table_type: int(core_table_counts_counter.get(table_type, 0)) for table_type in sorted(CORE_TABLE_TYPES)}
        has_all_core = all(core_table_counts.get(table_type, 0) >= 1 for table_type in CORE_TABLE_TYPES)
        has_exactly_one_each = all(core_table_counts.get(table_type, 0) == 1 for table_type in CORE_TABLE_TYPES)

        core_indexes_flat = [idx for indexes in core_table_indexes.values() for idx in indexes]
        last_core_table_index = max(core_indexes_flat) if core_indexes_flat else None
        last_core_table_type = table_type_by_index.get(last_core_table_index) if last_core_table_index is not None else None
        last_core_table_seq = next(
            (
                item.seq
                for item in items
                if item.kind == "table" and item.table_index == last_core_table_index
            ),
            None,
        )

        tail_items = [item for item in items if last_core_table_seq is not None and item.seq > last_core_table_seq]
        tail_table_types: Counter[str] = Counter()
        tail_contains_core_table = False
        for item in tail_items:
            if item.kind != "table" or item.table_index is None:
                continue
            table_type = table_type_by_index.get(item.table_index, "generic")
            tail_table_types[table_type] += 1
            if table_type in CORE_TABLE_TYPES:
                tail_contains_core_table = True

        tail_categories = Counter(_item_category(item) for item in tail_items)
        remaining_noise_items = sum(
            1 for item in items
            if any(marker in _norm(item.text) for marker in CONSULTANT_NOISE_MARKERS)
        )

        tail_samples = [
            {
                "seq": item.seq,
                "kind": item.kind,
                "table_index": item.table_index,
                "rows_count": item.rows_count,
                "category": _item_category(item),
                "table_type": table_type_by_index.get(item.table_index or -1) if item.kind == "table" else None,
                "text": _compact_sample(item.text),
            }
            for item in tail_items[:max_samples]
        ]

        if not has_all_core:
            warnings.append("not_all_core_tables_found_by_fast_heuristic")
        if not has_exactly_one_each:
            warnings.append("core_table_count_not_exactly_1_1_1_by_fast_heuristic")
        if tail_contains_core_table:
            warnings.append("tail_contains_core_table")
        if remaining_noise_items > 0:
            warnings.append("remaining_consultant_noise_items")

        return DocumentAnalysis(
            cleaned_filename=cleaned_path.name,
            raw_filename=raw_path.name if raw_path is not None else expected_raw_name,
            raw_found=raw_path is not None,
            raw_size_bytes=raw_size_bytes,
            raw_sha256=raw_sha256,
            cleaned_size_bytes=cleaned_size_bytes,
            cleaned_sha256=cleaned_sha256,
            raw_official_start_seq=raw_official_start,
            raw_items_before_official_start=raw_before_start,
            cleaned_official_start_seq=cleaned_official_start,
            cleaned_items_before_official_start=cleaned_before_start,
            cleaned_tables_count=len(cleaned_doc.tables),
            core_table_counts=core_table_counts,
            core_table_indexes=core_table_indexes,
            has_all_core_tables=has_all_core,
            has_exactly_one_each_core_table=has_exactly_one_each,
            last_core_table_index=last_core_table_index,
            last_core_table_type=last_core_table_type,
            last_core_table_seq=last_core_table_seq,
            trim_candidate=has_all_core and last_core_table_seq is not None and not tail_contains_core_table,
            tail_items_count=len(tail_items),
            tail_paragraphs_count=sum(1 for item in tail_items if item.kind == "paragraph"),
            tail_tables_count=sum(1 for item in tail_items if item.kind == "table"),
            tail_chars_count=sum(len(item.text) for item in tail_items),
            tail_table_type_counts=dict(sorted(tail_table_types.items())),
            tail_category_counts=dict(sorted(tail_categories.items())),
            tail_contains_core_table=tail_contains_core_table,
            tail_first_samples=tail_samples,
            remaining_consultant_noise_items=remaining_noise_items,
            warnings=warnings,
            error=None,
        )
    except Exception as exc:
        warnings.append("analysis_error")
        return DocumentAnalysis(
            cleaned_filename=cleaned_path.name,
            raw_filename=raw_path.name if raw_path is not None else expected_raw_name,
            raw_found=raw_path is not None,
            raw_size_bytes=raw_size_bytes,
            raw_sha256=raw_sha256,
            cleaned_size_bytes=cleaned_size_bytes,
            cleaned_sha256=cleaned_sha256,
            raw_official_start_seq=raw_official_start,
            raw_items_before_official_start=raw_before_start,
            cleaned_official_start_seq=None,
            cleaned_items_before_official_start=None,
            cleaned_tables_count=0,
            core_table_counts={table_type: 0 for table_type in sorted(CORE_TABLE_TYPES)},
            core_table_indexes={table_type: [] for table_type in sorted(CORE_TABLE_TYPES)},
            has_all_core_tables=False,
            has_exactly_one_each_core_table=False,
            last_core_table_index=None,
            last_core_table_type=None,
            last_core_table_seq=None,
            trim_candidate=False,
            tail_items_count=0,
            tail_paragraphs_count=0,
            tail_tables_count=0,
            tail_chars_count=0,
            tail_table_type_counts={},
            tail_category_counts={},
            tail_contains_core_table=False,
            tail_first_samples=[],
            remaining_consultant_noise_items=0,
            warnings=warnings,
            error=repr(exc),
        )


def _summarize(records: list[DocumentAnalysis]) -> RunSummary:
    core_patterns: Counter[str] = Counter()
    tail_table_types: Counter[str] = Counter()
    tail_categories: Counter[str] = Counter()
    for record in records:
        pattern = "/".join(str(record.core_table_counts.get(t, 0)) for t in ("identifiers", "documents", "refusal_reasons"))
        core_patterns[pattern] += 1
        tail_table_types.update(record.tail_table_type_counts)
        tail_categories.update(record.tail_category_counts)
    return RunSummary(
        analyzed_files=len(records),
        errors_count=sum(1 for r in records if r.error),
        raw_matched_files=sum(1 for r in records if r.raw_found),
        raw_official_start_found=sum(1 for r in records if r.raw_official_start_seq is not None),
        cleaned_official_start_found=sum(1 for r in records if r.cleaned_official_start_seq is not None),
        exact_core_1_1_1_files=sum(1 for r in records if r.has_exactly_one_each_core_table),
        all_core_found_files=sum(1 for r in records if r.has_all_core_tables),
        trim_candidate_files=sum(1 for r in records if r.trim_candidate),
        tail_items_total=sum(r.tail_items_count for r in records),
        tail_paragraphs_total=sum(r.tail_paragraphs_count for r in records),
        tail_tables_total=sum(r.tail_tables_count for r in records),
        tail_chars_total=sum(r.tail_chars_count for r in records),
        remaining_consultant_noise_items_total=sum(r.remaining_consultant_noise_items for r in records),
        tail_contains_core_table_files=sum(1 for r in records if r.tail_contains_core_table),
        core_count_patterns=dict(sorted(core_patterns.items())),
        tail_table_type_counts=dict(sorted(tail_table_types.items())),
        tail_category_counts=dict(sorted(tail_categories.items())),
    )


def _write_csv(records: list[DocumentAnalysis], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow([
            "cleaned_filename", "raw_filename", "raw_found",
            "raw_items_before_official_start", "cleaned_items_before_official_start",
            "cleaned_tables_count", "core_identifiers", "core_documents", "core_refusal_reasons",
            "core_table_indexes_json", "last_core_table_index", "last_core_table_type",
            "trim_candidate", "tail_items_count", "tail_paragraphs_count", "tail_tables_count",
            "tail_chars_count", "tail_table_type_counts_json", "tail_category_counts_json",
            "remaining_consultant_noise_items", "warnings_json", "error",
        ])
        for r in records:
            writer.writerow([
                r.cleaned_filename, r.raw_filename, r.raw_found,
                r.raw_items_before_official_start, r.cleaned_items_before_official_start,
                r.cleaned_tables_count, r.core_table_counts.get("identifiers", 0),
                r.core_table_counts.get("documents", 0), r.core_table_counts.get("refusal_reasons", 0),
                json.dumps(r.core_table_indexes, ensure_ascii=False), r.last_core_table_index,
                r.last_core_table_type, r.trim_candidate, r.tail_items_count, r.tail_paragraphs_count,
                r.tail_tables_count, r.tail_chars_count, json.dumps(r.tail_table_type_counts, ensure_ascii=False),
                json.dumps(r.tail_category_counts, ensure_ascii=False), r.remaining_consultant_noise_items,
                json.dumps(r.warnings, ensure_ascii=False), r.error,
            ])


def _write_review(summary: RunSummary, records: list[DocumentAnalysis], path: Path) -> None:
    lines: list[str] = []
    lines.append("АНАЛИЗ КАНДИДАТОВ НА ПРОГРАММНУЮ ОЧИСТКУ DOCX ПЕРЕД INGESTION")
    lines.append("")
    lines.append("ВАЖНО: это быстрый диагностический анализатор. Он ничего не меняет в DOCX и БД.")
    lines.append("Перед включением trim в production нужно сверить результаты с текущим extractor/QC после переингеста.")
    lines.append("")
    lines.append("ИТОГ")
    lines.append(f"- проанализировано cleaned DOCX: {summary.analyzed_files}")
    lines.append(f"- ошибок анализа: {summary.errors_count}")
    lines.append(f"- raw-файлов сопоставлено: {summary.raw_matched_files}")
    lines.append(f"- official start найден в raw: {summary.raw_official_start_found}")
    lines.append(f"- official start найден в cleaned: {summary.cleaned_official_start_found}")
    lines.append(f"- ключевые таблицы 1/1/1 по быстрой эвристике: {summary.exact_core_1_1_1_files}")
    lines.append(f"- все 3 ключевые таблицы найдены: {summary.all_core_found_files}")
    lines.append(f"- кандидатов на обрезку хвоста: {summary.trim_candidate_files}")
    lines.append(f"- файлов, где в хвосте есть ключевая таблица: {summary.tail_contains_core_table_files}")
    lines.append("")
    lines.append("ПОТЕНЦИАЛЬНО ОТРЕЗАЕМЫЙ ХВОСТ ПОСЛЕ ПОСЛЕДНЕЙ КЛЮЧЕВОЙ ТАБЛИЦЫ")
    lines.append(f"- элементов всего: {summary.tail_items_total}")
    lines.append(f"- абзацев: {summary.tail_paragraphs_total}")
    lines.append(f"- таблиц: {summary.tail_tables_total}")
    lines.append(f"- символов текста: {summary.tail_chars_total}")
    lines.append(f"- типы таблиц в хвосте: {json.dumps(summary.tail_table_type_counts, ensure_ascii=False)}")
    lines.append(f"- категории элементов хвоста: {json.dumps(summary.tail_category_counts, ensure_ascii=False)}")
    lines.append("")
    lines.append("РАСПРЕДЕЛЕНИЕ КЛЮЧЕВЫХ ТАБЛИЦ identifiers/documents/refusal_reasons")
    lines.append(json.dumps(summary.core_count_patterns, ensure_ascii=False, indent=2))
    lines.append("")

    problem_records = [r for r in records if r.error or not r.has_exactly_one_each_core_table or r.tail_contains_core_table]
    if problem_records:
        lines.append("ФАЙЛЫ ДЛЯ РУЧНОЙ ПРОВЕРКИ")
        for r in problem_records:
            lines.append(f"- {r.cleaned_filename}")
            lines.append(f"  warnings: {', '.join(r.warnings) or '-'}")
            lines.append(f"  core: {r.core_table_counts}, indexes: {r.core_table_indexes}")
            if r.error:
                lines.append(f"  error: {r.error}")
        lines.append("")

    top_tail = sorted(records, key=lambda r: r.tail_items_count, reverse=True)[:15]
    lines.append("ТОП-15 ФАЙЛОВ ПО РАЗМЕРУ ХВОСТА")
    for r in top_tail:
        lines.append(f"- {r.cleaned_filename}: items={r.tail_items_count}, tables={r.tail_tables_count}, chars={r.tail_chars_count}, last_core={r.last_core_table_index}/{r.last_core_table_type}")
        for sample in r.tail_first_samples[:3]:
            lines.append(f"  * {sample.get('kind')} #{sample.get('table_index') or '-'} {sample.get('category')}: {sample.get('text')}")
    lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def _run(args: argparse.Namespace) -> int:
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="rag2_docx_preprocess_analysis_") as tmp:
        tmp_dir = Path(tmp)
        cleaned_source = Path(args.cleaned_dir).resolve() if args.cleaned_dir else None
        if args.cleaned_zip:
            cleaned_source = tmp_dir / "cleaned"
            _safe_extract_zip(Path(args.cleaned_zip).resolve(), cleaned_source)
        if cleaned_source is None:
            raise ValueError("Need --cleaned-dir or --cleaned-zip")
        raw_source = Path(args.raw_dir).resolve() if args.raw_dir else None
        if args.raw_zip:
            raw_source = tmp_dir / "raw"
            _safe_extract_zip(Path(args.raw_zip).resolve(), raw_source)

        cleaned_files = _collect_docx_files(cleaned_source)
        if not cleaned_files:
            raise ValueError(f"No DOCX files found in cleaned source: {cleaned_source}")
        raw_files = _collect_docx_files(raw_source) if raw_source is not None and raw_source.exists() else []
        raw_map = _build_raw_map(raw_files)

        records: list[DocumentAnalysis] = []
        for index, cleaned_path in enumerate(cleaned_files, start=1):
            if not args.quiet:
                print(f"[{index}/{len(cleaned_files)}] {cleaned_path.name}")
            records.append(_analyze_one(cleaned_path, raw_map, args.max_samples))

        summary = _summarize(records)
        report_path = output_dir / "docx_preprocessing_candidates_report.json"
        summary_path = output_dir / "docx_preprocessing_candidates_summary.csv"
        review_path = output_dir / "docx_preprocessing_candidates_review.txt"
        report_path.write_text(json.dumps({"summary": asdict(summary), "records": [asdict(r) for r in records]}, ensure_ascii=False, indent=2), encoding="utf-8")
        _write_csv(records, summary_path)
        _write_review(summary, records, review_path)

        print("ГОТОВО")
        print(f"cleaned DOCX: {summary.analyzed_files}")
        print(f"ошибок анализа: {summary.errors_count}")
        print(f"raw сопоставлено: {summary.raw_matched_files}")
        print(f"ключевые таблицы 1/1/1 по быстрой эвристике: {summary.exact_core_1_1_1_files}")
        print(f"все 3 ключевые таблицы найдены: {summary.all_core_found_files}")
        print(f"кандидаты на trim: {summary.trim_candidate_files}")
        print(f"хвост после последней ключевой таблицы: {summary.tail_items_total} элементов, {summary.tail_tables_total} таблиц")
        print(f"типы таблиц в хвосте: {json.dumps(summary.tail_table_type_counts, ensure_ascii=False)}")
        print(f"категории хвоста: {json.dumps(summary.tail_category_counts, ensure_ascii=False)}")
        print(f"JSON: {report_path}")
        print(f"CSV:  {summary_path}")
        print(f"TXT:  {review_path}")
        return 1 if args.fail_on_warnings and summary.errors_count else 0


def _build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Диагностический анализ DOCX перед переносом очистки КонсультантПлюс из Word-макроса в ingestion."
    )
    parser.add_argument("--cleaned-dir", help="Папка с уже очищенными DOCX.")
    parser.add_argument("--cleaned-zip", help="ZIP с уже очищенными DOCX.")
    parser.add_argument("--raw-dir", help="Папка с raw DOCX из КонсультантПлюс.")
    parser.add_argument("--raw-zip", help="ZIP с raw DOCX из КонсультантПлюс.")
    parser.add_argument("--output-dir", default="logs/docx_preprocessing_analysis", help="Куда сохранить JSON/CSV/TXT отчёты.")
    parser.add_argument("--max-samples", type=int, default=8, help="Сколько первых элементов хвоста сохранить в JSON по каждому файлу.")
    parser.add_argument("--quiet", action="store_true", help="Не печатать имя каждого файла.")
    parser.add_argument("--fail-on-warnings", action="store_true", help="Завершаться с кодом 1 при ошибках анализа.")
    return parser


def main() -> int:
    parser = _build_arg_parser()
    args = parser.parse_args()
    if not args.cleaned_dir and not args.cleaned_zip:
        parser.error("Нужно указать --cleaned-dir или --cleaned-zip")
    if args.cleaned_dir and args.cleaned_zip:
        parser.error("Укажи только один источник cleaned: --cleaned-dir или --cleaned-zip")
    if args.raw_dir and args.raw_zip:
        parser.error("Укажи только один источник raw: --raw-dir или --raw-zip")
    return _run(args)


if __name__ == "__main__":
    raise SystemExit(main())
